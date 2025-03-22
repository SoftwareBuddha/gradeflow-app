
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from io import BytesIO
from PIL import Image

def add_formatted_page(doc, row):
    for col, val in row.dropna().items():
        if "Unnamed" not in col:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            run_label = p.add_run(f"{col}: ")
            run_label.bold = True
            run_label.font.size = Pt(11)

            run_value = p.add_run(str(val))
            run_value.font.size = Pt(11)
    doc.add_page_break()

st.set_page_config(page_title="GradeFlow - Excel/CSV to Word", layout="centered")

st.image("gf_logo.png", width=120)
st.title("📄 GradeFlow: Excel/CSV to Formatted Word")

st.markdown("Upload an Excel or CSV spreadsheet and generate a **Word document** where each row is formatted like a form on a separate page.")

uploaded_file = st.file_uploader("📤 Upload Excel or CSV File", type=["xlsx", "csv"])

if uploaded_file:
    try:
        if uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
        else:
            df = pd.read_excel(uploaded_file)

        # Clean the data
        df = df[~df.iloc[:, 0].astype(str).str.startswith(("Series", "Programme"))]
        df = df.dropna(how="all")
        df = df.loc[:, ~df.columns.str.contains('^Unnamed')]

        st.success("✅ File uploaded and cleaned!")
        st.write("Here’s a preview of the data that will be used:")
        st.dataframe(df)

        if st.button("🚀 Generate Word Document"):
            doc = Document()
            doc.add_heading("GradeFlow Generated Entries", 0)

            for _, row in df.iterrows():
                add_formatted_page(doc, row)

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                label="📥 Download Word Document",
                data=buffer,
                file_name="gradeflow_output.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"Something went wrong: {e}")
