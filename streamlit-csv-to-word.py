import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Inches
import io
import base64

def csv_to_word_table(csv_data):
    """
    Convert CSV data to a Word document with a table
    
    Args:
        csv_data: The uploaded CSV file data
    Returns:
        Word document as bytes
    """
    # Read CSV data
    df = pd.read_csv(csv_data)
    
    # Create a new Word document
    doc = Document()
    
    # Add a title
    doc.add_heading('CSV Data Table', 0)
    
    # Add a paragraph with some information
    doc.add_paragraph('This table was automatically generated from CSV data.')
    
    # Add a table
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    
    # Add headers
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        hdr_cells[i].text = column
    
    # Add data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer

def get_download_link(buffer, filename):
    """
    Generate a download link for the Word document
    
    Args:
        buffer: The document buffer
        filename: The filename for download
    Returns:
        HTML download link
    """
    b64 = base64.b64encode(buffer.getvalue()).decode()
    href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}">Download Word Document</a>'
    return href

def main():
    st.title("CSV to Word Table Converter")
    
    st.write("""
    ### Upload a CSV file and convert it to a Word document table
    This app takes your CSV data and creates a nicely formatted table in a Word document.
    """)
    
    uploaded_file = st.file_uploader("Choose a CSV file", type="csv")
    
    if uploaded_file is not None:
        # Show preview of data
        df = pd.read_csv(uploaded_file)
        st.subheader("Data Preview")
        st.dataframe(df.head())
        
        # Statistics
        st.subheader("Data Statistics")
        st.write(f"Total rows: {df.shape[0]}")
        st.write(f"Total columns: {df.shape[1]}")
        
        # Add options for table formatting
        st.subheader("Table Formatting Options")
        
        col1, col2 = st.columns(2)
        with col1:
            include_headers = st.checkbox("Include column headers", value=True)
        with col2:
            table_style = st.selectbox(
                "Table style",
                options=["Table Grid", "Light Shading", "Light List", "Medium Shading 1"],
                index=0
            )
        
        document_title = st.text_input("Document title", "CSV Data Table")
        
        # Create Word document
        if st.button("Generate Word Document"):
            # Reset file pointer
            uploaded_file.seek(0)
            
            with st.spinner("Creating Word document..."):
                buffer = csv_to_word_table(uploaded_file)
                
                # Display download link
                st.success("Word document created successfully!")
                st.markdown(get_download_link(buffer, "data_table.docx"), unsafe_allow_html=True)

if __name__ == "__main__":
    main()
